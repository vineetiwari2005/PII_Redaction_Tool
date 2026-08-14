"""
Glossary Extractor — auto-generates a deny-list from the document's
"Definitions and Abbreviations" section.

Every Red Herring Prospectus contains a glossary of defined terms
(financial, legal, technical). These are NOT PII but look like
organization or person names to statistical NER models. Extracting
them automatically keeps the deny-list up-to-date without manual
curation.
"""

import re
from typing import Set

from docx import Document


# Markers that signal the end of the definitions section
_SECTION_BOUNDARIES = {
    "forward-looking statements",
    "presentation of financial",
    "general information",
    "risk factors",
    "introduction",
    "summary of the offer",
    "the offer",
    "capital structure",
}


def extract_glossary_terms(doc: Document) -> Set[str]:
    """
    Walk through the document paragraphs and tables to collect
    terms defined in the glossary / abbreviations section.

    Handles three common formats:
      A. "TERM" means … (quoted terms)
      B. Bold term followed by its definition
      C. Two-column table (Term | Definition)

    Returns a set of lowercase, whitespace-normalised terms.
    """
    terms: Set[str] = set()

    inside_section = False

    for para in doc.paragraphs:
        text = para.text.strip()
        lo = text.lower()

        if not inside_section:
            if ("definition" in lo and "abbreviation" in lo) or lo in (
                "definitions and abbreviations",
                "definitions & abbreviations",
                "definitions",
                "abbreviations and definitions",
                "glossary",
                "glossary of terms",
            ):
                inside_section = True
            continue

        # Detect next section
        if any(marker in lo for marker in _SECTION_BOUNDARIES):
            if len(text) < 100:
                break

        if not text:
            continue

        # Format A: quoted terms
        for qt in re.findall(r'"([^"]+)"', text):
            if len(qt.split()) <= 6:
                terms.add(qt.strip().lower())

        # Format B: bold runs
        if para.runs:
            bold_chunk = ""
            for run in para.runs:
                if run.bold:
                    bold_chunk += run.text
                else:
                    break
            bold_chunk = bold_chunk.strip()
            if bold_chunk and len(bold_chunk.split()) <= 6:
                terms.add(bold_chunk.lower())

        # All-caps short headings
        if text.isupper() and len(text.split()) <= 6 and len(text) < 60:
            terms.add(text.lower())

    # Tables
    for table in doc.tables:
        if not table.rows:
            continue
        header = " ".join(c.text.strip().lower() for c in table.rows[0].cells)
        if not any(kw in header for kw in ("term", "abbreviation",
                                            "definition", "glossary", "acronym")):
            continue
        for row in table.rows[1:]:
            if row.cells:
                t = row.cells[0].text.strip()
                if t and len(t.split()) <= 6:
                    terms.add(t.lower())

    return {t for t in terms if len(t) > 2}


def extract_glossary_from_path(filepath: str) -> Set[str]:
    """Convenience wrapper: load a .docx and extract its glossary."""
    return extract_glossary_terms(Document(filepath))
