"""
DOCX Processor — reads and writes Word documents with run-level precision.

Handles the core challenge: PII detection operates on paragraph-level
concatenated text, but replacements must be applied to individual XML
runs to preserve bold, italic, font, colour, and other formatting.
"""

from dataclasses import dataclass
from typing import List, Optional, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
import os


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class RunSlice:
    """Tracks a single docx run's position within the concatenated paragraph text."""
    run: object           # docx.text.run.Run reference
    start: int            # start offset in concatenated text
    end: int              # end offset in concatenated text
    original_text: str    # text as it was before any edits


@dataclass
class DocSegment:
    """
    One paragraph of text together with its run-level offset map.

    Attributes:
        content:        Full concatenated text of all runs.
        slices:         Offset map back to individual runs.
        paragraph:      Reference to the docx Paragraph object.
        origin:         "body" for body paragraphs, "table" for table cells.
        origin_index:   Position tuple for traceability.
    """
    content: str
    slices: List[RunSlice]
    paragraph: Paragraph
    origin: str
    origin_index: tuple


# ---------------------------------------------------------------------------
# Segment extraction
# ---------------------------------------------------------------------------

def gather_segments(doc: Document) -> List[DocSegment]:
    """
    Walk through every paragraph and table cell in *doc*,
    returning a list of DocSegments with run-level offset maps.
    """
    segments: List[DocSegment] = []

    # Body paragraphs
    for idx, para in enumerate(doc.paragraphs):
        seg = _make_segment(para, origin="body", origin_index=(idx,))
        if seg:
            segments.append(seg)

    # Table cell paragraphs (with merged-cell deduplication)
    for t_idx, table in enumerate(doc.tables):
        visited: set = set()
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cid = id(cell._element)
                if cid in visited:
                    continue
                visited.add(cid)
                for p_idx, para in enumerate(cell.paragraphs):
                    seg = _make_segment(
                        para,
                        origin="table",
                        origin_index=(t_idx, r_idx, c_idx, p_idx),
                    )
                    if seg:
                        segments.append(seg)

    return segments


def _make_segment(
    para: Paragraph, origin: str, origin_index: tuple
) -> Optional[DocSegment]:
    """Build a DocSegment from a paragraph, or return None if it is empty."""
    runs = para.runs
    if not runs:
        return None

    slices: List[RunSlice] = []
    pos = 0
    for run in runs:
        txt = run.text
        slices.append(RunSlice(
            run=run, start=pos, end=pos + len(txt), original_text=txt,
        ))
        pos += len(txt)

    full = "".join(r.text for r in runs)
    if not full.strip():
        return None

    return DocSegment(
        content=full,
        slices=slices,
        paragraph=para,
        origin=origin,
        origin_index=origin_index,
    )


# ---------------------------------------------------------------------------
# Text replacement engine
# ---------------------------------------------------------------------------

def inject_replacements(
    segment: DocSegment,
    edits: List[Tuple[int, int, str]],
) -> None:
    """
    Replace text spans in a DocSegment while preserving run formatting.

    Args:
        segment: The segment whose runs will be modified in-place.
        edits:   List of (start, end, new_text) tuples relative
                 to segment.content.
    """
    if not edits:
        return

    edits = sorted(edits, key=lambda x: x[0])

    # Reconstruct the full text with replacements applied
    rebuilt = ""
    cursor = 0
    for s, e, replacement in edits:
        rebuilt += segment.content[cursor:s]
        rebuilt += replacement
        cursor = e
    rebuilt += segment.content[cursor:]

    # Map old character positions to new positions
    mapper = _offset_mapper(edits)

    # Distribute the new text back into each run
    for sl in segment.slices:
        ns = mapper(sl.start)
        ne = mapper(sl.end)
        ns = max(0, min(ns, len(rebuilt)))
        ne = max(ns, min(ne, len(rebuilt)))
        sl.run.text = rebuilt[ns:ne]


def _offset_mapper(edits: List[Tuple[int, int, str]]):
    """
    Returns a closure that translates old-text character positions
    into new-text positions after all edits have been applied.
    """
    intervals = []
    cumulative = 0
    for s, e, replacement in edits:
        span_len = e - s
        repl_len = len(replacement)
        intervals.append({
            "old_s": s,
            "old_e": e,
            "shift": cumulative,
            "repl_len": repl_len,
            "span_len": span_len,
        })
        cumulative += repl_len - span_len
    total = cumulative

    def translate(old_pos: int) -> int:
        for iv in intervals:
            if old_pos <= iv["old_s"]:
                return old_pos + iv["shift"]
            elif old_pos >= iv["old_e"]:
                continue
            else:
                # Inside a replacement — interpolate
                if iv["span_len"] > 0:
                    ratio = (old_pos - iv["old_s"]) / iv["span_len"]
                else:
                    ratio = 0.0
                base = iv["old_s"] + iv["shift"]
                return base + int(ratio * iv["repl_len"])
        return old_pos + total

    return translate


# ---------------------------------------------------------------------------
# Blackout formatting
# ---------------------------------------------------------------------------

def apply_blackout_style(segment: DocSegment, bar: str = "█") -> None:
    """
    Set black font colour and black highlight on every run that
    contains bar characters. Call AFTER inject_replacements().
    """
    from docx.shared import RGBColor
    from docx.enum.text import WD_COLOR_INDEX

    for sl in segment.slices:
        if bar in sl.run.text:
            sl.run.font.color.rgb = RGBColor(0, 0, 0)
            sl.run.font.highlight_color = WD_COLOR_INDEX.BLACK


# ---------------------------------------------------------------------------
# Persistence
# ---------------------------------------------------------------------------

def persist_document(doc: Document, path: str) -> None:
    """Save the modified document, creating parent directories as needed."""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
